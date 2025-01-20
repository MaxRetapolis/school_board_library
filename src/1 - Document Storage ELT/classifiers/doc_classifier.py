"""
doc_classifier.py

Contains classification logic for Microsoft Word documents:
- .doc (via a fallback or PyPandoc approach)
- .docx (via python-docx)
"""

import logging
from docx import Document as DocxDocument
import pypandoc

def classify_doc(filepath: str) -> dict:
    """
    Classify a .doc file by trying to extract text with pypandoc or a similar approach.
    
    Returns a dictionary with keys like "Text", "Images", "Tables".
    """
    # We'll reuse the docx classification logic if possible, or convert.
    # If you prefer a direct approach, handle .doc explicitly here.

    # Attempt: convert .doc to plain text for detection
    classification_results = {
        "Text": "No",
        "Images": "No",
        "Tables": "No"
    }

    try:
        # Convert .doc to plain text
        text_content = pypandoc.convert_file(filepath, 'plain')
        if text_content.strip():
            classification_results["Text"] = "Yes"
        # Checking images or tables in an older .doc can be trickier;
        # an alternative is to convert to docx temp file, then run docx logic.
        # ...
    except Exception as e:
        logging.error(f"Error classifying .doc file: {e}")
        classification_results["Text"] = "Unknown"

    return classification_results

def classify_docx(filepath: str) -> dict:
    """
    Classify a .docx file by checking for text, images, and tables using python-docx.
    """
    classification_results = {
        "Text": "No",
        "Images": "No",
        "Tables": "No"
    }

    try:
        doc = DocxDocument(filepath)

        # Check for text
        if any(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text):
            classification_results["Text"] = "Yes"

        # Check for images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                classification_results["Images"] = "Yes"
                break

        # Check for tables
        if doc.tables:
            classification_results["Tables"] = "Yes"

    except Exception as e:
        logging.error(f"Error classifying .docx file: {e}")
        classification_results["Text"] = "Unknown"
        classification_results["Images"] = "Unknown"
        classification_results["Tables"] = "Unknown"

    return classification_results
